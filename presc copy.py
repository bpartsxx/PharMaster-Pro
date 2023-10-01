from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image

def get_run_formatting(run):
    return {
        "font_size": run.font.size,
        "font_name": run.font.name,
        "bold": run.bold,
        "italic": run.italic,
        # Add more formatting attributes as needed
    }

try:
    # Load the template document
    template_doc = Document("prescription.docx")  # Replace with your template file

    # Define data to replace placeholders
    data = {
        "[PATIENT_NAME]": "John Doe",
        "[DOCTOR_NAME]": "Dr. Smith",
        "[PRESCRIPTION_DATE]": "2023-09-13",
        "[list_med]": ["Item 1", "Item 2", "Item 3"],  # Add a list as a placeholder
    }

    # Replace text placeholders while preserving formatting
    for p in template_doc.paragraphs:
        for key, value in data.items():
            if key in p.text:
                if isinstance(value, list):  # Check if the value is a list
                    # Remove the paragraph containing the list placeholder
                    template_doc.paragraphs.remove(p)
                    # Insert the list items
                    for item in value:
                        new_paragraph = template_doc.add_paragraph(item)
                        new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Set alignment as needed
                else:
                    for run in p.runs:
                        if key in run.text:
                            original_format = get_run_formatting(run)
                            run.text = run.text.replace(key, value)
                            # Apply the original formatting to the new text
                            new_run = p.runs[-1]  # Get the last run with the replaced text
                            new_run.bold = original_format["bold"]
                            new_run.italic = original_format["italic"]
                            new_run.font.size = original_format["font_size"]
                            new_run.font.name = original_format["font_name"]
                            # Add more formatting attributes as needed

    # Insert an image
    image_path = "qr.png"  # Replace with your image file
    img = Image.open(image_path)
    width, height = img.size
    if width > 4.5 * 72:  # Resize image if it's wider than 4.5 inches (assuming 72 DPI)
        img.thumbnail((4.5 * 72, height))
    template_doc.add_picture(image_path, width=Pt(img.width), height=Pt(img.height))

    # Save the final document
    template_doc.save("prescription1.docx")  # Replace with the desired output filename
except Exception as e:
    print(f"An error occurred: {e}")