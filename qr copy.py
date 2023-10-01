import mysql.connector
import qrcode
import png
from pyqrcode import QRCode
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import pdfkit
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.lib import colorsissue
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet
from PIL import Image

last_processed_id=3

def get_run_formatting(run):
    return {
        "font_size": run.font.size,
        "font_name": run.font.name,
        "bold": run.bold,
        "italic": run.italic,
        # Add more formatting attributes as needed
    }


while True:
    
    con=mysql.connector.connect(host="localhost", user="root", password="", database="pms_db", port=3310)
    cursor = con.cursor()
    medicines=[]
   
    query1 = "SELECT pmh.patient_visit_id, p.id AS patient_id, p.patient_name AS patient_name, p.phone_number AS patient_phone, m.medicine_name, pmh.quantity, pmh.dosage, pv.disease, p.date_of_birth FROM patient_visits pv JOIN patient_medication_history pmh ON pv.id = pmh.patient_visit_id JOIN medicine_details md ON pmh.medicine_details_id = md.id JOIN medicines m ON md.medicine_id = m.id JOIN patients p ON pv.patient_id = p.id WHERE pmh.patient_visit_id > %s;"

    cursor.execute(query1,(last_processed_id,))
            
    # display all records
    table = cursor.fetchall()
    if len(table)!=0:
        name=table[0][2]
        dose=table[0][6]
        disease=table[0][7]
        dob=str(table[0][8])
        if len(table[0][3])==10:
            number=table[0][3]
        else:
            print("invalid number")
            break
        for row in table:
            medicines.append(row[4]+"  x"+str(row[5])+" "+str(row[6])+"\n")
        last_processed_id=table[0][0]
        
        outfile="qr.png"
        qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_L, box_size=10, border=4)
        for med in medicines:
            qr.add_data(med)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")


        img.save(outfile)

        medicinestring='\n'.join(medicines)
        try:
    # Load the template document
            template_doc = Document("prescription.docx")  # Replace with your template file
            
            # Define data to replace placeholders
            data = {
                "[name]": name,
                "[age]":dob,
                "[doctor_name]": "Dr. Smith",
                "[number]":number,
                "[med]":medicinestring,
                "[disease]":disease,
                "image_path": "qr.png"
                
            }

            # Replace text placeholders while preserving formatting
            for p in template_doc.paragraphs:
                for key, value in data.items():
                    if key in p.text:
                        for run in p.runs:
                            if key in run.text:
                                #original_format = get_run_formatting(run)
                                run.text = run.text.replace(key, value)
                                # Apply the original formatting to the new text
                                new_run = p.runs[-1]  # Get the last run with the replaced text
                                #new_run.bold = original_format["bold"]
                                #new_run.italic = original_format["italic"]
                                #new_run.font.size = original_format["font_size"]
                                #new_run.font.name = original_format["font_name"]
                                # Add more formatting attributes as needed

            # Insert an image
            image_path = "qr.png"  # Replace with your image file
            img = Image.open(image_path)
            max_width = Inches(0.5)
            max_height = Inches(0.5)

            img.thumbnail((max_width, max_height))
            template_doc.add_picture(image_path, width=Pt(img.width), height=Pt(img.height))

            # Save the final document
            template_doc.save("prescription1.docx")  # Replace with the desired output filename
        except Exception as e:
            print(f"An error occurred: {e}")


        #wkhtmltopdf_path = r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe'

        #input_file = "prescription1.docx"  # Replace with your input DOCX file name
        #output_file = "prescription.pdf"  # Replace with the desired output PDF file name
        pdf_filename = "output.pdf"
        doc = SimpleDocTemplate(pdf_filename, pagesize=letter)
        story = []

        title = Paragraph("Prescription", getSampleStyleSheet()["Title"])
        story.append(title)
        story.append(Spacer(1, 12))

        # Add content from your data
        content = f"Doctor Name: {data['[doctor_name]']}<br/>"
        content += f"Patient Name: {data['[name]']}<br/>"
        content += f"Patient dob: {data['[age]']}<br/>"
        content += f"Phone Number: {data['[number]']}<br/>"
        content += f"Medication: {data['[med]']}<br/>"
        content += f"Disease: {data['[disease]']}"    




        image = Image(image_path, width=200, height=100)  # Set the width and height as desired

        # Append the image to the story
        story.append(image)

        # Optionally, add spacing after the image
        spacer = Spacer(1, 20)  # Adjust the height (20 in this example) as needed
        story.append(spacer)

        # Build the PDF document
        doc.build(story)


    cursor.close()
    con.close()
    time.sleep(5)

    print(last_processed_id)