import mysql.connector
import qrcode
import png
from pyqrcode import QRCode
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import pdfkit
from docx.shared import Inches
from docx2pdf import convert
from twilio.rest import Client
import os
import socket

#ACc8663f990403b4e3364ccce00c86b0c4
#1ccad53fcf426fa0001ba4991aeeef6d
#+12569739503

account_sid = 'ACc8663f990403b4e3364ccce00c86b0c4'
auth_token = '1ccad53fcf426fa0001ba4991aeeef6d'

client = Client(account_sid, auth_token)

last_processed_id=1

def get_run_formatting(run):
    return {
        "font_size": run.font.size,
        "font_name": run.font.name,
        "bold": run.bold,
        "italic": run.italic,
        # Add more formatting attributes as needed
    }


while True:
    
    con=mysql.connector.connect(host="localhost", user="root", password="", database="pms_db", port=3310)
    cursor = con.cursor()
    medicines=[]
   
    query1 = "SELECT pmh.patient_visit_id, p.id AS patient_id, p.patient_name AS patient_name, p.phone_number AS patient_phone, m.medicine_name, pmh.quantity, pmh.dosage, pv.disease, p.date_of_birth FROM patient_visits pv JOIN patient_medication_history pmh ON pv.id = pmh.patient_visit_id JOIN medicine_details md ON pmh.medicine_details_id = md.id JOIN medicines m ON md.medicine_id = m.id JOIN patients p ON pv.patient_id = p.id WHERE pmh.patient_visit_id > %s;"

    cursor.execute(query1,(last_processed_id,))
            
    # display all records
    table = cursor.fetchall()
    if len(table)!=0:
        name=table[0][2]
        dose=table[0][6]
        disease=table[0][7]
        dob=str(table[0][8])
        
        number=table[0][3]
        for row in table:
            medicines.append(row[4]+"  x"+str(row[5])+" "+str(row[6])+"\n")
        last_processed_id=table[0][0]
        
        outfile="qr.png"
        qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_L, box_size=10, border=4)
        for med in medicines:
            qr.add_data(med)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")


        img.save(outfile)

        medicinestring='\n'.join(medicines)
        try:
    # Load the template document
            template_doc = Document("prescription.docx")  # Replace with your template file
            
            # Define data to replace placeholders
            data = {
                "[name]": name,
                "[age]":dob,
                "[doctor_name]": "Dr. Smith",
                "[number]":number,
                "[med]":medicinestring,
                "[disease]":disease
                
            }

            # Replace text placeholders while preserving formatting
            for p in template_doc.paragraphs:
                for key, value in data.items():
                    if key in p.text:
                        for run in p.runs:
                            if key in run.text:
                                #original_format = get_run_formatting(run)
                                run.text = run.text.replace(key, value)
                                # Apply the original formatting to the new text
                                new_run = p.runs[-1]  # Get the last run with the replaced text
                                #new_run.bold = original_format["bold"]
                                #new_run.italic = original_format["italic"]
                                #new_run.font.size = original_format["font_size"]
                                #new_run.font.name = original_format["font_name"]
                                # Add more formatting attributes as needed

            # Insert an image
            input_image = Image.open('qr.png')

            # Define the new dimensions (width, height)
            new_width = 120  # Replace with your desired width
            new_height = 120  # Replace with your desired height

            # Resize the image
            resized_image = input_image.resize((new_width, new_height))

            # Save the resized image to a new file
            resized_image.save('qr.png')

            # Close the original and resized images (optional, but good practice)
            input_image.close()
            resized_image.close()

            image_path = "qr.png"  # Replace with your image file
            img = Image.open(image_path)
            max_width = Inches(0.1)
            max_height = Inches(0.1)

            img.thumbnail((max_width, max_height))
            template_doc.add_picture(image_path, width=Pt(img.width), height=Pt(img.height))

            # Save the final document
            template_doc.save("prescription1.docx")  # Replace with the desired output filename
        except Exception as e:
            print(f"An error occurred: {e}")


        os.system('server.pyw')
        input_file = "prescription1.docx"  # Replace with your input DOCX file name
        output_file = "prescription.pdf"  # Replace with the desired output PDF file name
        
        
        convert(input_file, output_file)
        #hostname = socket.gethostname()
        #ip = socket.gethostbyname(hostname)
        #file_url = 'http://localhost:8000/prescription.pdf'
        #message = client.messages \
         #       .create(
          #           body="Join Earth's mightiest heroes. Like Kevin Bacon.",
           #          from_='+12569739503',
            #         to='+919474027940'
             #    )
        #print('whatsapp:+91'+number)
        #print(message.sid)

        print(last_processed_id)



    cursor.close()
    con.close()
    time.sleep(3)

       